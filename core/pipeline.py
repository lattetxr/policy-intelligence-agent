"""三条业务流水线编排：政策解析 → 政企匹配 → 落地评估报告。

模块 1  政策PDF解析：policy_docs/*.pdf → 切片 → RAG向量库 → 通义千问提取要素
模块 2  政企匹配：企业经营指标 → 向量检索 → 通义千问打分 + 申报优先级
模块 3  落地评估：科研团队反馈台账 → 通义千问生成标准化评估文档 → output/
"""
import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

import pandas as pd
from loguru import logger

from .schemas import (
    ExtractedPolicy, EnterpriseNeed, MatchResult,
    FeedbackEntry, FeedbackLedger, PolicyAssessment, ApplicationStandard,
)
from .tools import PDFParser, RAGRetriever, PolicyExtractor, PolicyMatcher
from .models import ModelClient, build_messages

DEFAULT_MODEL = "tongyi"


class PolicyPipeline:
    def __init__(
        self,
        model_client: ModelClient,
        retriever: RAGRetriever,
        extractor: PolicyExtractor,
        matcher: PolicyMatcher,
    ):
        self.model = model_client
        self.retriever = retriever
        self.extractor = extractor
        self.matcher = matcher

    # ───────────────────────────────────────────────────────
    # 模块 1：政策PDF解析
    # policy_docs/*.pdf → 切片 → RAG → 千问提取 → data/processed/*.json
    # ───────────────────────────────────────────────────────
    def ingest_policy_pdfs(
        self,
        pdf_dir: str = "policy_docs",
        model_key: str = DEFAULT_MODEL,
    ) -> Dict[str, Any]:
        pdfs = PDFParser.list_pdfs(pdf_dir)
        if not pdfs:
            return {"ok": False, "message": f"{pdf_dir} 下没有PDF文件", "processed": []}

        parser = PDFParser()
        ingested, extracted_list = [], []
        for path in pdfs:
            doc_id = parser.file_hash(path)
            text, chunks = parser.process_pdf(path)
            if not chunks:
                logger.warning(f"{path.name} 解析后为空，跳过")
                continue

            n = self.retriever.add_documents(doc_id, chunks)
            if n > 0:
                ingested.append({"file": path.name, "doc_id": doc_id, "chunks": n})
                logger.info(f"[模块1] 已入库 {path.name}: {n} 个切片")

        if not ingested:
            return {"ok": True, "message": "PDF已全部入库（或重复导入）", "processed": []}

        logger.info(f"[模块1] 开始调用 {model_key} 提取政策要素 ...")
        for item in ingested:
            path = Path(pdf_dir) / item["file"]
            text, _ = parser.process_pdf(path)
            policy = self.extractor.extract(text, title=path.stem, model_key=model_key)
            policy.source_file = item["file"]
            policy.id = item["doc_id"]
            self._save_processed_policy(policy)
            extracted_list.append(policy)

        return {
            "ok": True,
            "message": f"完成 {len(ingested)} 个PDF入库与要素提取",
            "processed": ingested,
            "policies": extracted_list,
        }

    def _save_processed_policy(self, policy: ExtractedPolicy) -> str:
        out_dir = Path("data/processed")
        out_dir.mkdir(parents=True, exist_ok=True)
        safe = policy.title.replace("/", "_").replace(" ", "_")[:50] or "未命名"
        path = out_dir / f"{safe}.json"
        path.write_text(json.dumps(policy.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
        return str(path)

    def load_processed_policies(self, folder: str = "data/processed") -> List[ExtractedPolicy]:
        policies = []
        for fp in sorted(Path(folder).glob("*.json")):
            try:
                data = json.loads(fp.read_text(encoding="utf-8"))
                policies.append(ExtractedPolicy(**data))
            except Exception as e:
                logger.warning(f"加载 {fp.name} 失败: {e}")
        return policies

    def get_system_stats(self) -> Dict[str, Any]:
        """首页仪表盘统计：入库政策 / 匹配服务 / 材料清单 / 评估报告 四维计数"""
        def count_dir(patterns):
            total = 0
            for pat in patterns:
                total += len(list(Path(".").glob(pat)))
            return total

        from core.tools import PDFParser
        return {
            "policy_total": len(self.load_processed_policies()),
            "pdf_total": len(PDFParser.list_pdfs("policy_docs")),
            "chunks_total": self.retriever.collection.count() if hasattr(self, "retriever") else 0,
            "match_services": count_dir(["output/matching/*.json", "output/enterprise_match/*.json"]),
            "standard_plans": count_dir(["output/application/*.md"]),
            "assessment_reports": count_dir(["output/assessment/*.md"]),
        }

    # ───────────────────────────────────────────────────────
    # 模块 2：政企匹配
    # 企业经营指标 → 向量检索相关政策 → 千问打分+优先级 → 匹配报告
    # ───────────────────────────────────────────────────────
    def match_enterprise(
        self,
        enterprise: EnterpriseNeed,
        top_k: int = 5,
        model_key: str = DEFAULT_MODEL,
    ) -> Dict[str, Any]:
        if self.retriever.collection.count() == 0:
            return {"ok": False, "message": "向量库为空，请先运行模块1导入政策PDF", "matches": []}

        query = self._build_enterprise_query(enterprise)
        retrieved = self.retriever.search(query, top_k=top_k)

        policies = self._retrieved_to_policies(retrieved)
        if not policies:
            return {"ok": False, "message": "向量检索未命中政策", "matches": []}

        logger.info(f"[模块2] 命中 {len(policies)} 个政策，调用 {model_key} 匹配打分 ...")
        matches = self.matcher.batch_match(enterprise, policies, model_key)

        self._save_matching_json(enterprise, matches)
        return {
            "ok": True,
            "message": f"完成 {len(matches)} 个政策匹配",
            "matches": matches,
            "ranked": self._rank_priority(matches),
        }

    def _build_enterprise_query(self, e: EnterpriseNeed) -> str:
        parts = [e.industry, e.size]
        if e.revenue:
            parts.append(f"营收{e.revenue}万元")
        if e.rnd_ratio:
            parts.append(f"研发占比{e.rnd_ratio}")
        if e.patent_count:
            parts.append(f"专利{e.patent_count}项")
        parts += e.needs
        return " ".join([p for p in parts if p]) or "企业政策需求"

    def _retrieved_to_policies(self, retrieved: List[Dict[str, Any]]) -> List[ExtractedPolicy]:
        processed = {p.id: p for p in self.load_processed_policies()}
        policies, seen = [], set()
        for item in retrieved:
            meta = item.get("metadata", {})
            doc_id = meta.get("doc_id", "")
            if doc_id in seen:
                continue
            seen.add(doc_id)
            if doc_id in processed:
                policies.append(processed[doc_id])
                continue
            from .schemas import ExtractedPolicy
            policies.append(ExtractedPolicy(
                id=doc_id,
                title=meta.get("title", doc_id[:8]),
                raw_text_snippet=item["text"],
            ))
        return policies

    @staticmethod
    def _rank_priority(matches: List[MatchResult]) -> List[Dict[str, Any]]:
        order = {"高": 0, "中": 1, "低": 2}
        ranked = sorted(matches, key=lambda m: (-m.match_score, order.get(m.priority_level, 3)))
        return [
            {"rank": i + 1, "policy_title": m.policy_title, "match_score": m.match_score,
             "priority_level": m.priority_level, "priority_reason": m.priority_reason}
            for i, m in enumerate(ranked)
        ]

    def _save_matching_json(self, enterprise: EnterpriseNeed, matches: List[MatchResult]) -> str:
        out_dir = Path("output/matching")
        out_dir.mkdir(parents=True, exist_ok=True)
        safe = enterprise.name.replace("/", "_").replace(" ", "_")[:50] or "未知企业"
        path = out_dir / f"{safe}_匹配明细_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        payload = {
            "enterprise": enterprise.model_dump(),
            "generated_at": datetime.now().isoformat(),
            "ranked": self._rank_priority(matches),
            "matches": [m.model_dump() for m in matches],
        }
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(path)

    # ───────────────────────────────────────────────────────
    # 模块 3：落地评估报告
    # 科研团队反馈台账 → 千问生成标准化评估文档 → output/assessment/
    # ───────────────────────────────────────────────────────
    def generate_assessment_report(
        self,
        ledger: FeedbackLedger,
        model_key: str = DEFAULT_MODEL,
    ) -> Dict[str, Any]:
        if not ledger.feedback_items:
            return {"ok": False, "message": "反馈台账为空", "report_path": ""}

        context = self._build_report_context(ledger)

        ASSESS_PROMPT = """你是政务政策落地评估专家。基于科研团队反馈台账和政策信息，生成一份标准化《政策落地评估报告》。

【政策背景】
{context}

请按以下结构生成报告正文（Markdown格式，中文）：

## 一、政策落地总体情况
（结合反馈数量、解决率、主要问题领域概述）

## 二、主要问题与原因分析
（按问题分类归纳，逐类分析根因：申报流程、材料要求、政策解读、资金到位等）

## 三、阶段性成效评估
（列举已解决/进行中事项及价值）

## 四、风险与障碍提示
（未解决问题对政策落地的潜在影响）

## 五、优化建议
（针对流程、材料、解读、资金四条线给出可操作建议）

## 六、下一阶段工作安排
（需求收集-问题分析-优化迭代的闭环机制建议）"""

        try:
            response = self.model.chat(
                model_key,
                build_messages("你是政务评估专家，输出规范的Markdown报告正文。", ASSESS_PROMPT.format(context=context)),
                temperature=0.3,
                max_tokens=4096,
            )
            report_path = self._save_assessment_ledger(ledger, response, model_key)
            return {"ok": True, "message": "评估报告已生成", "report_path": report_path, "content": response}
        except Exception as e:
            logger.error(f"[模块3] 报告生成失败: {e}")
            return {"ok": False, "message": f"报告生成失败: {e}", "report_path": ""}

    def _build_report_context(self, ledger: FeedbackLedger) -> str:
        lines = [f"政策名称：{ledger.policy_title}", f"科研团队：{ledger.research_team}", f"评估周期：{ledger.period}"]
        if self.retriever.collection.count() > 0 and ledger.policy_title:
            hits = self.retriever.search(ledger.policy_title, top_k=3)
            if hits:
                lines.append("\n【政策原文摘要】")
                lines.extend(f"- {h['text'][:200]}" for h in hits)

        lines.append(f"\n【反馈台账明细】共 {len(ledger.feedback_items)} 条：")
        for i, item in enumerate(ledger.feedback_items, 1):
            lines.append(
                f"{i}. [{item.date}] {item.team} | 分类: {item.category} | 状态: {item.status} | "
                f"内容: {item.content} | 影响: {item.impact or '—'}"
            )
        return "\n".join(lines)

    def _save_assessment_ledger(self, ledger: FeedbackLedger, content: str, model_key: str) -> str:
        from .tools import ReportGenerator
        from .schemas import ExtractedPolicy

        policy = ExtractedPolicy(title=ledger.policy_title or "政策落地评估", source_file="反馈台账")
        path = ReportGenerator.save_assessment_report(policy, content)
        meta = Path(path).with_suffix(".json")
        meta.write_text(json.dumps({
            "ledger": ledger.model_dump(), "model": model_key, "generated_at": datetime.now().isoformat(),
        }, ensure_ascii=False, indent=2), encoding="utf-8")
        return path

    # ───────────────────────────────────────────────────────
    # 模块4（迭代1）：政策申报材料标准化
    # 针对专项政府项目申报：申报全链路流程 + 材料规范要点 + 审核易错点 + 标准化材料清单
    # ───────────────────────────────────────────────────────
    def generate_application_standard(
        self,
        policy: Optional[ExtractedPolicy] = None,
        model_key: str = DEFAULT_MODEL,
    ) -> Dict[str, Any]:
        if policy is None:
            policies = self.load_processed_policies()
            if not policies:
                return {"ok": False, "message": "暂无已提取政策，请先运行模块1", "standard": None}
            policy = policies[0]

        process_items = [
            item["text"] for item in self._retrieve_policy_chunks(policy.id)
        ]
        context = "\n\n".join(
            [f"[政策:{policy.title}]\n{policy.raw_text_snippet[:1500]}"] + process_items
        )

        STANDARD_PROMPT = """你是政务专项项目申报标准化专家，服务政府/事业单位/科研实验室的专项项目申报场景。
基于给定政策原文，梳理申报全链路流程、材料规范要点、审核易错点，并输出标准化申报材料清单。输出严格JSON（不要Markdown代码块）。

【政策原文】
{context}

输出JSON结构（字段名必须精确）：
{{
  "application_process": [
    {{"step": "执行环节名", "content": "关键内容", "responsible": "责任主体", "duration": "预计周期"}}
  ],
  "material_points": ["材料规范要点1（含格式/盖章/份数等硬性要求）", "材料规范要点2"],
  "audit_mistakes": ["审核易错点1（含规避建议）", "审核易错点2"],
  "material_checklist": [
    {{"item": "材料名称", "format": "格式要求", "content_points": "内容要点", "tips": "提交节点/注意事项"}}
  ],
  "quality_control": ["质量管控清单项1", "质量管控清单项2"]
}}

要求：
- application_process 按时间顺序覆盖申报准备→线上填报→材料递交→初审→复审/专家评审→公示→拨付立项等节点，6-10步
- material_checklist 覆盖申报书/预算表/财务报表/知识产权证明/合作协议/承诺书等，5-8条
- audit_mistakes 3-6条，均须基于政策原文，不得编造"""

        try:
            response = self.model.chat(
                model_key,
                build_messages("你是政务申报标准化专家，只输出合法JSON。", STANDARD_PROMPT.format(context=context)),
                temperature=0.3,
                max_tokens=4096,
                json_mode=True,
            )
            data = json.loads(response)
            standard = ApplicationStandard(
                policy_title=policy.title, policy_id=policy.id,
                model_used=model_key, **data,
            )
            md_path = self._save_application_standard(policy, standard, model_key)
            return {"ok": True, "message": "申报材料标准化方案已生成",
                    "path": md_path, "standard": standard,
                    "content": self._render_application_markdown(standard)}
        except Exception as e:
            logger.error(f"[模块4] 申报材料标准化生成失败: {e}")
            return {"ok": False, "message": f"申报材料标准化生成失败: {e}", "standard": None}

    @staticmethod
    def _render_application_markdown(s: "ApplicationStandard") -> str:
        md = [f"# {s.policy_title} · 申报材料标准化方案", "", f"- 生成时间：{s.generated_at}　模型：{s.model_used}", "",
              "## 一、申报全链路流程", "", "| 执行环节 | 关键内容 | 责任主体 | 预计周期 |", "| --- | --- | --- | --- |"]
        for p in s.application_process:
            md.append(f"| {p.get('step','')} | {p.get('content','')} | {p.get('responsible','')} | {p.get('duration','')} |")
        md += ["", "## 二、材料规范要点", ""]
        md += [f"{i+1}. {x}" for i, x in enumerate(s.material_points)]
        md += ["", "## 三、审核易错点", ""]
        md += [f"{i+1}. {x}" for i, x in enumerate(s.audit_mistakes)]
        md += ["", "## 四、标准化申报材料清单", "",
               "| 材料名称 | 格式要求 | 内容要点 | 提交节点/注意事项 |", "| --- | --- | --- | --- |"]
        for c in s.material_checklist:
            md.append(f"| {c.get('item','')} | {c.get('format','')} | {c.get('content_points','')} | {c.get('tips','')} |")
        md += ["", "## 五、质量管控清单", ""]
        md += [f"{i+1}. {x}" for i, x in enumerate(s.quality_control)]
        return "\n".join(md)

    def _save_application_standard(self, policy: ExtractedPolicy, s: "ApplicationStandard", model_key: str) -> str:
        out_dir = Path("output/application")
        out_dir.mkdir(parents=True, exist_ok=True)
        safe = re.sub(r"[\\/:*?\"<>|\s]+", "_", policy.title) or "policy"
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = out_dir / f"{safe}_申报材料标准化_{stamp}.md"
        path.write_text(self._render_application_markdown(s), encoding="utf-8")
        logger.info(f"[模块4] 申报材料标准化已保存: {path}")
        meta = path.with_suffix(".json")
        meta.write_text(json.dumps(s.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
        return str(path)

    def _retrieve_policy_chunks(self, doc_id: str, top_k: int = 5) -> List[Dict[str, Any]]:
        return self.retriever.get_by_doc_id(doc_id, top_k=top_k)

    # ───────────────────────────────────────────────────────
    # 模块5（迭代2）：政策-企业需求匹配分析
    # 输入企业服务痛点 + 经营维度 → 检索本地政策库 → 匹配度分析 + 适配政策条目 + 落地建议
    # ───────────────────────────────────────────────────────
    def analyze_enterprise_match(
        self,
        pain_points: str,
        dimensions: Optional[Dict[str, Any]] = None,
        top_k: int = 5,
        model_key: str = DEFAULT_MODEL,
    ) -> Dict[str, Any]:
        if self.retriever.collection.count() == 0:
            return {"ok": False, "message": "向量库为空，请先运行模块1导入政策PDF", "report_path": ""}

        dimensions = dimensions or {}
        dim_lines = [f"- {k}：{v}" for k, v in dimensions.items() if v]
        query = f"{pain_points} " + " ".join(f"{k}{v}" for k, v in dimensions.items() if v)
        retrieved = self.retriever.search(query, top_k=top_k)
        policies = self._retrieved_to_policies(retrieved)
        if not policies:
            return {"ok": False, "message": "向量检索未命中政策", "report_path": ""}

        policy_lines = []
        for p in policies:
            policy_lines.append(
                f"【{p.title}】\n申报门槛：{'；'.join(p.eligibility_criteria) or '未知'}\n"
                f"适用主体：{'；'.join(p.applicable_entities) or '未知'}\n"
                f"补贴标准：{'；'.join(f'{s.subsidy_type}:{s.amount_or_ratio}' for s in p.subsidy_standards) or '未知'}\n"
                f"量化阈值：{'；'.join(f'{k}:{v}' for k, v in p.policy_thresholds.items()) or '未知'}"
            )
        policy_text = "\n\n".join(policy_lines)

        ANALYZE_PROMPT = """你是政府扶持政策-中小企业需求匹配专家，服务于政府/园区/科研实验室的政策落地与企业服务工作。
基于中小企业的服务痛点和经营维度，从本地政策库中匹配最适配的扶持政策，输出结构化匹配分析报告（Markdown格式，中文）。

【企业服务痛点】
{pain_points}

【企业经营维度】
{dim_lines}

【本地政策库（已检索命中）】
{policy_text}

请按以下结构输出《企业需求-政策匹配分析报告》：

## 一、企业画像与服务痛点解读
（归纳企业经营维度，提炼核心服务痛点与政策诉求方向）

## 二、政策匹配度分析
（逐政策分析匹配度，给出匹配度评分（0-100）与核心理由，说明政策对该企业痛点/经营维度的支撑点）

## 三、适配政策条目
（从命中的政策中列出具体可申报的扶持条目：政策名称→适用条件→具体扶持内容/补贴标准→申报时间窗口）

## 四、落地建议
（按优先级排序给出可操作建议：优先申报哪个政策、需补齐哪些条件/材料、预期收益、风险提示）

报告结尾给出"结论"一句话总结推荐申报路径。"""

        try:
            response = self.model.chat(
                model_key,
                build_messages("你是企业政策匹配专家，输出规范的Markdown匹配分析报告。", ANALYZE_PROMPT.format(
                    pain_points=pain_points,
                    dim_lines="\n".join(dim_lines) or "未提供，请基于痛点推断",
                    policy_text=policy_text,
                )),
                temperature=0.3,
                max_tokens=4096,
            )
            path = self._save_enterprise_match_report(pain_points, dimensions, policies, response, model_key)
            return {"ok": True, "message": f"完成 {len(policies)} 个政策的需求匹配分析",
                    "report_path": path, "content": response, "matched_policies": [p.title for p in policies]}
        except Exception as e:
            logger.error(f"[模块5] 需求匹配分析失败: {e}")
            return {"ok": False, "message": f"需求匹配分析失败: {e}", "report_path": ""}

    def _save_enterprise_match_report(
        self, pain_points: str, dimensions: Dict[str, Any],
        policies: List[ExtractedPolicy], content: str, model_key: str,
    ) -> str:
        out_dir = Path("output/enterprise_match")
        out_dir.mkdir(parents=True, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        name = dimensions.get("name") or "企业需求匹配"
        safe = re.sub(r"[\\/:*?\"<>|\s]+", "_", name) or "enterprise"
        path = out_dir / f"{safe}_需求匹配分析_{stamp}.md"
        header = (
            f"# {name} · 企业需求-政策匹配分析报告\n\n"
            f"- 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}　模型：{model_key}\n"
            f"- 命中政策：{('、'.join(p.title for p in policies))}\n\n---\n\n"
        )
        path.write_text(header + content, encoding="utf-8")
        logger.info(f"[模块5] 需求匹配分析已保存: {path}")
        meta = path.with_suffix(".json")
        meta.write_text(json.dumps({
            "pain_points": pain_points, "dimensions": dimensions,
            "matched_policies": [p.title for p in policies],
            "model": model_key, "generated_at": datetime.now().isoformat(),
        }, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(path)

    # ───────────────────────────────────────────────────────
    # 模块6（迭代3）：政策落地评估报告自动生成
    # 政策条文 + 适用人群 + 落地约束 + 用户反馈 → 结构化优势/落地难点/迭代建议 → 可保存文档
    # ───────────────────────────────────────────────────────
    def generate_policy_assessment(
        self,
        policy: Optional[ExtractedPolicy] = None,
        feedback_text: str = "",
        model_key: str = DEFAULT_MODEL,
    ) -> Dict[str, Any]:
        if policy is None:
            policies = self.load_processed_policies()
            if not policies:
                return {"ok": False, "message": "暂无已提取政策，请先运行模块1", "assessment": None}
            policy = policies[0]

        clause_chunks = [item["text"] for item in self._retrieve_policy_chunks(policy.id)]
        feedback = feedback_text.strip() or "未提供用户反馈台账。"
        clauses_part = ("【政策原文片段】\n" + "\n".join(clause_chunks)) if clause_chunks else ""
        context = (
            f"【政策条文】\n{policy.raw_text_snippet[:1200]}\n\n"
            f"{clauses_part}\n\n"
            f"【适用人群】\n适用主体：{'；'.join(policy.applicable_entities) or '未知'}\n"
            f"申报门槛：{'；'.join(policy.eligibility_criteria) or '未知'}\n\n"
            f"【落地约束】\n约束条款：{'；'.join(policy.constraint_clauses) or '未知'}\n"
            f"量化阈值：{'；'.join(f'{k}:{v}' for k, v in policy.policy_thresholds.items()) or '未知'}\n\n"
            f"【用户反馈】\n{feedback}"
        )

        ASSESS_JSON_PROMPT = """你是政务政策落地评估专家。基于给定维度，对政策落地效果进行评估，输出严格JSON（不要Markdown代码块）。

【输入维度】
{context}

输出JSON结构（字段名必须精确）：
{{
  "overall": "总体评估结论（100字以内，概括政策设计质量与落地成效）",
  "strengths": [{{"title": "优势要点名", "detail": "具体说明（含条文/数据支撑）"}}],
  "landing_difficulties": [{{"title": "落地难点名", "detail": "具体说明（含涉及条款/主体）"}}],
  "feedback_analysis": "用户反馈维度分析（结合反馈台账，归纳共性问题、解决率与痛因）",
  "iteration_suggestions": ["迭代建议1（可操作、针对性强）", "迭代建议2"]
}}

要求：strengths 与 landing_difficulties 各 3-5 条；每条 detail 需结合政策条文原文，不得编造。"""

        try:
            raw = self.model.chat(
                model_key,
                build_messages("你是政务政策评估专家，只输出合法JSON。", ASSESS_JSON_PROMPT.format(context=context)),
                temperature=0.3,
                max_tokens=4096,
                json_mode=True,
            )
            data = json.loads(raw)
            assessment = PolicyAssessment(
                policy_title=policy.title, policy_id=policy.id,
                model_used=model_key, **data,
            )
            md_path, docx_path = self._save_policy_assessment(assessment)
            return {"ok": True, "message": "政策落地评估报告已生成", "assessment": assessment,
                    "md_path": md_path, "docx_path": docx_path}
        except Exception as e:
            logger.error(f"[模块6] 政策落地评估报告生成失败: {e}")
            return {"ok": False, "message": f"政策落地评估报告生成失败: {e}", "assessment": None}

    def _save_policy_assessment(self, a: PolicyAssessment):
        out_dir = Path("output/assessment")
        out_dir.mkdir(parents=True, exist_ok=True)
        safe = re.sub(r"[\\/:*?\"<>|\s]+", "_", a.policy_title) or "policy"
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        md = [
            f"# {a.policy_title} · 政策落地效果评估报告",
            "",
            f"- 评估时间：{a.generated_at}　模型：{a.model_used}",
            f"- 政策ID：{a.policy_id}",
            "",
            "---", "",
            "## 一、总体评估结论", a.overall, "",
            "## 二、政策优势", "",
        ]
        for s in a.strengths:
            md += [f"### {s.get('title','')}", s.get("detail", ""), ""]
        md += ["## 三、落地难点", ""]
        for d in a.landing_difficulties:
            md += [f"### {d.get('title','')}", d.get("detail", ""), ""]
        md += ["## 四、用户反馈维度分析", a.feedback_analysis, "",
               "## 五、迭代建议", ""]
        md += [f"{i+1}. {s}" for i, s in enumerate(a.iteration_suggestions)]
        md += ["", "---", f"报告由 {a.model_used} 自动生成，仅供参考，最终以主管部门解释为准。"]
        md_text = "\n".join(md)

        md_path = out_dir / f"{safe}_落地评估_{stamp}.md"
        md_path.write_text(md_text, encoding="utf-8")

        docx_path = out_dir / f"{safe}_落地评估_{stamp}.docx"
        try:
            self._render_assessment_docx(a, docx_path)
        except Exception as e:
            logger.warning(f"[模块6] docx 渲染失败（已保留md）: {e}")
            docx_path = None

        logger.info(f"[模块6] 评估报告已保存: {md_path}" + (f" / {docx_path}" if docx_path else ""))
        meta = out_dir / f"{safe}_落地评估_{stamp}.json"
        meta.write_text(json.dumps(a.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
        return str(md_path), (str(docx_path) if docx_path else "")

    @staticmethod
    def _render_assessment_docx(a: "PolicyAssessment", path: Path) -> None:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document()
        doc.add_heading(f"{a.policy_title} · 政策落地效果评估报告", level=0)
        doc.add_paragraph(f"评估时间：{a.generated_at}　模型：{a.model_used}")

        def section(title, items):
            doc.add_heading(title, level=1)
            for it in items:
                if isinstance(it, dict):
                    doc.add_heading(it.get("title", ""), level=2)
                    p = doc.add_paragraph(it.get("detail", ""))
                else:
                    doc.add_paragraph(it, style="List Number" if title == "五、迭代建议" else None)

        section("一、总体评估结论", [a.overall])
        section("二、政策优势", a.strengths)
        section("三、落地难点", a.landing_difficulties)
        doc.add_heading("四、用户反馈维度分析", level=1)
        doc.add_paragraph(a.feedback_analysis)
        section("五、迭代建议", a.iteration_suggestions)
        doc.add_paragraph()
        doc.add_paragraph("报告由 %s 自动生成，仅供参考，最终以主管部门解释为准。" % a.model_used)
        doc.save(str(path))

    # ───────────────────────────────────────────────────────
    # 反馈台账加载（输入入口：data/input/feedback/*.json | *.xlsx | *.csv）
    # ───────────────────────────────────────────────────────
    @staticmethod
    def load_feedback_ledger(path: str) -> FeedbackLedger:
        p = Path(path)
        if p.suffix.lower() in (".xlsx", ".xls"):
            return PolicyPipeline._ledger_from_excel(p)
        if p.suffix.lower() == ".csv":
            return PolicyPipeline._ledger_from_csv(p)
        return PolicyPipeline._ledger_from_json(p)

    @staticmethod
    def _ledger_from_json(p: Path) -> FeedbackLedger:
        data = json.loads(p.read_text(encoding="utf-8"))
        if "feedback_items" in data:
            return FeedbackLedger(**data)
        return FeedbackLedger(
            policy_title=data.get("policy_title", ""),
            research_team=data.get("research_team", ""),
            period=data.get("period", ""),
            feedback_items=[FeedbackEntry(**item) for item in data.get("items", [])],
        )

    @staticmethod
    def _ledger_from_excel(p: Path) -> FeedbackLedger:
        df = pd.read_excel(p)
        meta = {}
        for col in df.columns:
            if "政策" in str(col) and "名称" in str(col):
                meta["policy_title"] = str(df[col].dropna().iloc[0])
            elif "团队" in str(col):
                meta["research_team"] = str(df[col].dropna().iloc[0])
        return PolicyPipeline._ledger_from_dataframe(df, meta)

    @staticmethod
    def _ledger_from_csv(p: Path) -> FeedbackLedger:
        df = pd.read_csv(p)
        return PolicyPipeline._ledger_from_dataframe(df, {})

    @staticmethod
    def _ledger_from_dataframe(df, meta: Dict[str, str]) -> FeedbackLedger:
        items = []
        for _, row in df.iterrows():
            items.append(FeedbackEntry(
                date=str(row.get("日期", "") or ""),
                team=str(row.get("团队", "") or ""),
                category=str(row.get("分类", "") or ""),
                content=str(row.get("内容", "") or ""),
                status=str(row.get("状态", "") or ""),
                impact=str(row.get("影响", "") or ""),
            ))
        return FeedbackLedger(
            policy_title=meta.get("policy_title", "未命名政策"),
            research_team=meta.get("research_team", "科研团队"),
            period=meta.get("period", ""),
            feedback_items=items,
        )


def list_feedback_files(folder: str = "data/input/feedback") -> List[str]:
    Path(folder).mkdir(parents=True, exist_ok=True)
    return [str(p) for p in Path(folder).glob("*") if p.suffix.lower() in (".json", ".xlsx", ".xls", ".csv")]
